from datetime import datetime
from docx import Document
import requests
from urllib.parse import urlparse


class UserDocument:
    def __init__(self, doc_type, number, given_by, issue_date):
        self.type = doc_type
        self.number = number
        self.given_by = given_by
        self.issue_date = issue_date


class User:
    def __init__(
        self,
        name,
        email,
        phone,
        birth_date,
        doc_type,
        number,
        given_by,
        issue_date,
        address,
        degree,
        program,
        direction,
    ):
        self.name = name
        self.email = email
        self.phone = phone
        self.birth_date = birth_date
        self.document = UserDocument(doc_type, number, given_by, issue_date)
        self.address = address
        self.degree = degree
        self.program = program
        self.direction = direction


class Customer:
    def __init__(
        self,
        payment_type,
        name=None,
        address=None,
        doc_type=None,
        number=None,
        given_by=None,
        phone=None,
        document=None,
        document_date=None,
        ogrn=None,
        inn=None,
        behalf=None,
        behalf_genitive=None,
    ):
        self.payment_type = payment_type
        if payment_type == "individual":
            self.name = name
            self.document = UserDocument(doc_type, number, given_by, issue_date=None)
            self.phone = phone
            self.address = address
        elif payment_type == "juridical":
            self.address = address
            self.name = name
            self.document = (document,)
            self.document_date = document_date
            self.ogrn = ogrn
            self.inn = inn
            self.behalf = behalf
            self.behalf_genitive = behalf_genitive


def get_program_by_financing(selected_programs):
    idx = 0
    for program in selected_programs:
        if program["financing"] == "contract":
            return idx, program["program"]
        idx += 1
    return None, None


def get_data(cookies, current_url):
    parsed_url = urlparse(current_url)
    path_parts = parsed_url.path.split("/")
    user_id = path_parts[4]
    details_id = path_parts[5]

    application = requests.get(
        f"https://abitlk.itmo.ru/api/v1/users/{user_id}/studentDetails/" f"{details_id}/application",
        headers={"Cookie": cookies},
    ).json()
    personally = requests.get(
        f"https://abitlk.itmo.ru/api/v1/users/{user_id}/studentDetails/" f"{details_id}/forms/personally",
        headers={"Cookie": cookies},
    ).json()
    address = requests.get(
        f"https://abitlk.itmo.ru/api/v1/users/{user_id}/studentDetails/" f"{details_id}/forms/address",
        headers={"Cookie": cookies},
    ).json()
    programs = requests.get(
        f"https://abitlk.itmo.ru/api/v1/users/{user_id}/studentDetails/" f"{details_id}/programs",
        headers={"Cookie": cookies},
    ).json()
    payment = requests.get(
        f"https://abitlk.itmo.ru/api/v1/users/{user_id}/studentDetails/" f"{details_id}/forms/payment",
        headers={"Cookie": cookies},
    ).json()

    for response in (application, personally, address, programs, payment):
        if not response["ok"]:
            raise ConnectionError(response["message"])

    contract_idx, contract_program = get_program_by_financing(application["result"]["selected_programs"])

    user = User(
        name=application["result"]["user"]["full_name"],
        email=application["result"]["user"]["email"],
        phone=application["result"]["user"]["phone"],
        birth_date=application["result"]["user"]["birth_date"],
        doc_type=personally["result"]["documents"]["person_id"]["type"],
        number=application["result"]["user"]["passport_number"],
        given_by=personally["result"]["documents"]["person_id"]["division_name"],
        issue_date=personally["result"]["documents"]["person_id"]["issued_date"],
        address=address["result"]["registration_address"],
        degree=application["result"]["user"]["degree"],
        program=contract_program,
        direction=programs["result"]["selected_programs"][contract_idx]["competitive_group"]["title"]
        if application["result"]["user"]["degree"] == "bachelor"
        else programs["result"]["selected_programs"][contract_idx]["program"]["direction_of_education"],
    )
    user.birth_date = datetime.fromisoformat(user.birth_date.replace("Z", "+00:00"))
    user.document.issue_date = datetime.fromisoformat(user.document.issue_date.replace("Z", "+00:00"))

    if payment["result"]["payment_type"] == "individual":
        customer = Customer(
            payment_type=payment["result"]["payment_type"],
            name=payment["result"]["individual"]["full_name"]["full_name"],
            address=payment["result"]["individual"]["address"],
            doc_type="Паспорт гражданина РФ",
            number=payment["result"]["individual"]["series"] + payment["result"]["individual"]["number"],
            given_by=payment["result"]["individual"]["division_name"],
            phone=payment["result"]["individual"]["phone"],
        )
    elif payment["result"]["payment_type"] == "juridical":
        customer = Customer(
            payment_type=payment["result"]["payment_type"],
            name=payment["result"]["juridical"]["name"],
            address=payment["result"]["juridical"]["address"],
            document=payment["result"]["juridical"]["document_name_genitive"],
            document_date=payment["result"]["juridical"]["date_document"],
            ogrn=payment["result"]["juridical"]["ogrn"],
            inn=payment["result"]["juridical"]["inn"],
            behalf=payment["result"]["juridical"]["full_name_behalf"],
            behalf_genitive=payment["result"]["juridical"]["full_name_behalf_genitive"],
        )
        customer.document_date = datetime.fromisoformat(customer.document_date.replace("Z", "+00:00"))
    else:
        customer = Customer(payment_type=payment["result"]["payment_type"])

    return user, customer


def fill_contract(user, customer, contract_path="../logic/data/sample_contract_api.docx"):
    annual = {
        "Бизнес-информатика": "301000 (Триста одна тысяча)",
        "Технологии и инновации": "301000 (Триста одна тысяча)",
        "Управление высокотехнологичным бизнесом": "349000 (Триста сорок девять тысяч)",
        "Технологии и стратегии бизнес-трансформации": "349000 (Триста сорок девять тысяч)",
        "Цифровые продукты: создание и управление": "380000 (Триста восемьдесят тысяч)",
        "Стратегическое управление интеллектуальной собственностью / IP Management Strategy": "349000 (Триста сорок "
        "девять тысяч)",
    }
    total = {
        "Бизнес-информатика": "1204000 (Один миллион двести четыре тысячи)",
        "Технологии и инновации": "1204000 (Один миллион двести четыре тысячи)",
        "Управление высокотехнологичным бизнесом": "698000 (Шестьсот девяносто восемь тысяч)",
        "Технологии и стратегии бизнес-трансформации": "698000 (Шестьсот девяносто восемь тысяч)",
        "Цифровые продукты: создание и управление": "760000 (Семьсот шестьдесят тысяч)",
        "Стратегическое управление интеллектуальной собственностью / IP Management Strategy": "698000 (Шестьсот "
        "девяносто восемь тысяч)",
    }

    contract = Document(contract_path)

    today = datetime.today()
    user_age = (
        today.year - user.birth_date.year - ((today.month, today.day) < (user.birth_date.month, user.birth_date.day))
    )
    end_flag = False

    for paragraph in contract.paragraphs:
        for run in paragraph.runs:
            if "DATE" in run.text:
                months = {7: "июля", 8: "августа"}
                d, m, y = today.day, today.month, today.year
                run.text = run.text.replace("DATE", f"«{d}» {months[m]} {y}")
            if "CUSTOMER" in run.text:
                if customer.payment_type == "self":
                    run.text = run.text.replace("CUSTOMER", user.name)
                else:
                    run.text = run.text.replace("CUSTOMER", customer.name)
            if "FACE" in run.text:
                if customer.payment_type == "juridical":
                    run.text = run.text.replace("FACE", customer.behalf_genitive)
                else:
                    run.text = run.text.replace("FACE", "-")
            if "DOCUMENT" in run.text:
                if customer.payment_type == "juridical":
                    run.text = run.text.replace(
                        "DOCUMENT",
                        f" {customer.document} от " f"{customer.document_date.strftime('%d.%m.%Y')}",
                    )
                else:
                    run.text = run.text.replace("DOCUMENT", "")
            if "STUDENTA" in run.text:
                if customer.payment_type == "self":
                    run.text = run.text.replace("STUDENTA", "он же")
                else:
                    run.text = run.text.replace("STUDENTA", user.name)
            if "FORM" in run.text:
                if user.degree == "bachelor":
                    run.text = run.text.replace("FORM", "бакалавриата")
                else:
                    run.text = run.text.replace("FORM", "магистратуры")
            if "DIRECTION" in run.text:
                run.text = run.text.replace("DIRECTION", user.direction)
            if "PROGRAM" in run.text:
                run.text = run.text.replace("PROGRAM", "«" + user.program + "»")
            if "PERIOD" in run.text:
                if user.degree == "bachelor":
                    run.text = run.text.replace("PERIOD", "4")
                else:
                    run.text = run.text.replace("PERIOD", "2")
            if "DEGREE" in run.text:
                if user.degree == "bachelor":
                    run.text = run.text.replace("DEGREE", "бакалавра")
                else:
                    run.text = run.text.replace("DEGREE", "магистра")
            if "TOTAL" in run.text:
                run.text = run.text.replace("TOTAL", total[user.program])
            if "ANNUAL" in run.text:
                run.text = run.text.replace("ANNUAL", annual[user.program])
            if "PASSPORTCUST" in run.text:
                if customer.payment_type == "juridical":
                    run.text = run.text.replace(
                        "PASSPORTCUST",
                        "адрес: {0}, ОГРН {1}, ИНН {2}\n\n{3}".format(
                            customer.address, customer.ogrn, customer.inn, customer.behalf
                        ),
                    )
                elif customer.payment_type == "individual":
                    run.text = run.text.replace(
                        "PASSPORTCUST",
                        "Паспорт гражданина РФ {0}, выдан {1}, адрес: {2}, "
                        "{3}\n\n{4}".format(
                            customer.document.number,
                            customer.document.given_by,
                            customer.address,
                            customer.phone,
                            customer.name,
                        ),
                    )
                else:
                    run.text = run.text.replace(
                        "PASSPORTCUST",
                        "Паспорт гражданина РФ {0}, выдан {1} {2}, адрес: {3}, "
                        "{4}, {5}\n\n{6}".format(
                            user.document.number,
                            user.document.given_by,
                            user.document.issue_date.strftime("%d.%m.%Y"),
                            user.address,
                            user.phone,
                            user.email,
                            user.name,
                        ),
                    )
            if "STUDENTB" in run.text:
                if customer.payment_type == "self":
                    run.text = run.text.replace("STUDENTB", "См. графу «ЗАКАЗЧИК»")
                else:
                    run.text = run.text.replace(
                        "STUDENTB",
                        "{0}; Паспорт гражданина РФ {1}, выдан {2} {3}, адрес: "
                        "{4}, {5}, {6}\n\n{7}".format(
                            user.name,
                            user.document.number,
                            user.document.given_by,
                            user.document.issue_date.strftime("%d.%m.%Y"),
                            user.address,
                            user.phone,
                            user.email,
                            user.name,
                        ),
                    )
            if "Отметка о согласии на заключение настоящего Договора" in run.text:
                end_flag = True
            if user_age >= 18 and end_flag:
                run.text = ""

    contract.save("/".join(contract_path.split("/")[:-1]) + "/output_contract.docx")


def fill_receipt(user, customer, semesters, receipt_path="../logic/data/sample_receipt.docx"):
    annual_amount = {
        "Бизнес-информатика": 301000,
        "Технологии и инновации": 301000,
        "Управление высокотехнологичным бизнесом": 349000,
        "Технологии и стратегии бизнес-трансформации": 349000,
        "Цифровые продукты: создание и управление": 380000,
        "Стратегическое управление интеллектуальной собственностью / IP Management Strategy": 349000,
    }
    receipt = Document(receipt_path)

    for row in receipt.tables[0].rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if "CUSTOMER" in run.text:
                        if customer.payment_type == "self":
                            run.text = run.text.replace("CUSTOMER", user.name)
                        else:
                            run.text = run.text.replace("CUSTOMER", customer.name)
                    if "STUDENT" in run.text:
                        run.text = run.text.replace("STUDENT", user.name)
                    if "SEMESTER" in run.text:
                        if int(semesters) == 1:
                            run.text = run.text.replace("SEMESTER", "1 семестр")
                        else:
                            run.text = run.text.replace("SEMESTER", "1,2 семестры")
                    if "SUM" in run.text:
                        if int(semesters) == 1:
                            run.text = run.text.replace("SUM", str(int(0.55 * annual_amount[user.program])))
                        else:
                            run.text = run.text.replace("SUM", str(annual_amount[user.program]))

    receipt.save("/".join(receipt_path.split("/")[:-1]) + "/output_receipt.docx")


def autofill(cookies, current_url, semesters):
    user, customer = get_data(cookies, current_url)
    fill_contract(user, customer)
    fill_receipt(user, customer, semesters)
